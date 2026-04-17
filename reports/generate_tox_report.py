"""
Generate Toxicological Report for CALQIX OralBiome Pro — English + German (.docx)
Run: python reports/generate_tox_report.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, datetime

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

def shade(cell, hex_color):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn('w:shd'),{qn('w:fill'):hex_color,qn('w:color'):'auto',qn('w:val'):'clear'}))

def tbl(doc, hdrs, rows):
    t = doc.add_table(rows=1, cols=len(hdrs)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(hdrs):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9); r.font.name='Calibri'; r.font.color.rgb=RGBColor(255,255,255); shade(c,'1A3B5C')
    for rd in rows:
        row=t.add_row()
        for i,v in enumerate(rd):
            c=row.cells[i]; c.text=''; r=c.paragraphs[0].add_run(str(v)); r.font.size=Pt(9); r.font.name='Calibri'
    return t

def h1(doc,txt): doc.add_heading(txt,level=1)
def h2(doc,txt): doc.add_heading(txt,level=2)
def h3(doc,txt): doc.add_heading(txt,level=3)
def p(doc,txt,bold=False,italic=False):
    pr=doc.add_paragraph(); rn=pr.add_run(txt); rn.bold=bold; rn.italic=italic; return pr
def field(doc,label,val):
    pr=doc.add_paragraph(); rn=pr.add_run(label+': '); rn.bold=True; pr.add_run(val)

def hazard_block(doc, title, data, labels):
    h3(doc, title)
    for key, lbl in labels:
        if key in data:
            pr = doc.add_paragraph()
            rn = pr.add_run(lbl + ': '); rn.bold = True; rn.font.size = Pt(10)
            rn2 = pr.add_run(data[key]); rn2.font.size = Pt(10)

def build(lang):
    doc = Document()
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
    st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10)
    today = datetime.date.today().strftime('%d %B %Y')
    EN = lang=='en'

    # ── Cover ──
    for _ in range(4): doc.add_paragraph()
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn=cp.add_run('TOXICOLOGICAL ASSESSMENT REPORT' if EN else 'TOXIKOLOGISCHES BEWERTUNGSGUTACHTEN'); rn.bold=True; rn.font.size=Pt(22); rn.font.color.rgb=RGBColor(26,59,92)
    doc.add_paragraph()
    cp2=doc.add_paragraph(); cp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rn2=cp2.add_run('OralBiome Pro — '+('Oral Probiotic Lozenge' if EN else 'Orale Probiotische Lutschtablette')); rn2.font.size=Pt(14); rn2.font.color.rgb=RGBColor(26,59,92)
    doc.add_paragraph()
    cp3=doc.add_paragraph(); cp3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cp3.add_run(('Toxicological Safety Assessment of an Oral Probiotic Food Supplement' if EN else 'Toxikologische Sicherheitsbewertung eines oralen probiotischen Nahrungserganzungsmittels')).font.size=Pt(11)
    for _ in range(6): doc.add_paragraph()
    cp4=doc.add_paragraph(); cp4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rn4=cp4.add_run('CONFIDENTIAL' if EN else 'VERTRAULICH'); rn4.bold=True; rn4.font.size=Pt(12); rn4.font.color.rgb=RGBColor(192,0,0)
    cn=doc.add_paragraph(); cn.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cn.add_run(('This document contains proprietary and confidential information. Reproduction, distribution, or disclosure to third parties without prior written consent of the commissioning party is prohibited.' if EN else 'Dieses Dokument enthalt geschutzte und vertrauliche Informationen. Die Vervielfaltigung, Weitergabe oder Offenlegung an Dritte ohne vorherige schriftliche Genehmigung des Auftraggebers ist untersagt.')).font.size=Pt(9)
    doc.add_page_break()

    # ── 1. General Info ──
    h1(doc, '1. '+('General Information' if EN else 'Allgemeine Angaben'))
    gen_fields = [
        ('Report Title' if EN else 'Berichtstitel', ('Toxicological Safety Assessment — OralBiome Pro Oral Probiotic Lozenge' if EN else 'Toxikologische Sicherheitsbewertung — OralBiome Pro Orale Probiotische Lutschtablette')),
        ('Report Number' if EN else 'Berichtsnummer', 'TOX-CALQIX-OBP-[XXXX]-001'),
        ('Version', '1.0 — '+('Draft for Review' if EN else 'Entwurf zur Uberprufung')),
        ('Date' if EN else 'Datum', '[DD Month YYYY]' if EN else '[TT Monat JJJJ]'),
        ('Commissioner' if EN else 'Auftraggeber', 'CALQIX B.V.'),
        ('Author / Toxicologist' if EN else 'Verfasser / Toxikologe', '[Name], [Degree(s)]\n[Title / Function]\n[Registration number]' if EN else '[Name], [Akademische Grade]\n[Titel / Funktion]\n[Registrierungsnummer]'),
        ('Company Details' if EN else 'Unternehmensdaten', 'CALQIX B.V.\n[Address]\nThe Netherlands\nKVK: [number]\nVAT: [number]' if EN else 'CALQIX B.V.\n[Adresse]\nNiederlande\nKvK: [Nummer]\nUSt-IdNr.: [Nummer]'),
        ('Contact' if EN else 'Kontakt', 'Email: info@calqix.com\nTel: [number]' if EN else 'E-Mail: info@calqix.com\nTel: [Nummer]'),
    ]
    for lbl,val in gen_fields: field(doc,lbl,val)

    # ── 2. Purpose ──
    h1(doc, '2. '+('Purpose of the Assessment' if EN else 'Zweck der Bewertung'))
    if EN:
        p(doc, 'The purpose of this report is to provide a toxicological safety assessment of the product "OralBiome Pro" (hereinafter: "the Product"), an oral probiotic food supplement in dissolvable lozenge form, manufactured and marketed by CALQIX B.V.')
        p(doc, 'This assessment evaluates whether the Product, when used as directed by the adult consumer population, presents an acceptable level of safety with respect to its qualitative and quantitative composition, the intended route of exposure, and the anticipated frequency and duration of use.')
        p(doc, 'The scope encompasses:\n- Identification and characterisation of all ingredients\n- Hazard identification based on available toxicological data\n- Exposure assessment under reasonably foreseeable conditions of use\n- Risk characterisation including calculation of Margins of Safety (MoS) / Margins of Exposure (MoE)\n- Identification of data gaps, uncertainties, and limitations')
        p(doc, 'Regulatory context: This assessment takes into account Regulation (EC) No 178/2002 (General Food Law), Regulation (EC) No 1924/2006 (Nutrition and Health Claims), Directive 2002/46/EC (Food Supplements), and relevant national implementing legislation.')
    else:
        p(doc, 'Zweck dieses Berichts ist die toxikologische Sicherheitsbewertung des Produkts "OralBiome Pro" (nachfolgend: "das Produkt"), eines oralen probiotischen Nahrungserganzungsmittels in Form einer auflosbaren Lutschtablette, hergestellt und vertrieben von CALQIX B.V.')
        p(doc, 'Diese Bewertung pruft, ob das Produkt bei bestimmungsgemasser Verwendung durch die erwachsene Verbraucherpopulation ein akzeptables Sicherheitsniveau aufweist.')
        p(doc, 'Der Umfang umfasst:\n- Identifizierung und Charakterisierung aller Inhaltsstoffe\n- Gefahrenidentifizierung auf Grundlage verfugbarer toxikologischer Daten\n- Expositionsabschatzung unter vernunftigerweise vorhersehbaren Verwendungsbedingungen\n- Risikocharakterisierung einschliesslich Berechnung von Sicherheitsmargen (MoS) / Expositionsmargen (MoE)\n- Identifizierung von Datenlucken, Unsicherheiten und Einschrankungen')
        p(doc, 'Regulatorischer Kontext: Verordnung (EG) Nr. 178/2002, Verordnung (EG) Nr. 1924/2006, Richtlinie 2002/46/EG und einschlagige nationale Umsetzungsgesetzgebung.')

    # ── 3. Product ID ──
    h1(doc, '3. '+('Product Identification' if EN else 'Produktidentifizierung'))
    prod_fields = [
        ('Product Name' if EN else 'Produktname', 'OralBiome Pro'),
        ('Product Type' if EN else 'Produkttyp', ('Food supplement (oral probiotic)' if EN else 'Nahrungserganzungsmittel (orales Probiotikum)')),
        ('Dosage Form' if EN else 'Darreichungsform', ('Dissolvable lozenge (tablet for buccal/sublingual dissolution)' if EN else 'Auflosbare Lutschtablette (Tablette zur bukkalen/sublingualen Auflosung)')),
        ('Description' if EN else 'Beschreibung', ('OralBiome Pro is a daily oral probiotic food supplement containing four clinically studied bacterial strains (Lactobacillus reuteri DSM 17938, Lactobacillus paracasei, Streptococcus salivarius K12 (BLIS K12), Streptococcus salivarius M18 (BLIS M18)), combined with xylitol and natural peppermint flavouring, formulated as a dissolvable lozenge.' if EN else 'OralBiome Pro ist ein tagliches orales probiotisches Nahrungserganzungsmittel mit vier klinisch untersuchten Bakterienstammen (Lactobacillus reuteri DSM 17938, Lactobacillus paracasei, Streptococcus salivarius K12, Streptococcus salivarius M18), kombiniert mit Xylitol und naturlichem Pfefferminzaroma.')),
        ('Batch / Reference' if EN else 'Charge / Referenz', '[Batch number(s)]' if EN else '[Chargennummer(n)]'),
        ('Manufacturer' if EN else 'Hersteller', '[Name and address]\n[GMP/HACCP/ISO certification]' if EN else '[Name und Adresse]\n[GMP/HACCP/ISO-Zertifizierung]'),
        ('Target Population' if EN else 'Zielpopulation', ('Adults (18+). Not intended for children who cannot safely dissolve a lozenge.' if EN else 'Erwachsene (ab 18 Jahren). Nicht bestimmt fur Kinder.')),
        ('Intended Use' if EN else 'Vorgesehene Verwendung', ('Daily oral care supplement. 1 lozenge dissolved slowly in the mouth after brushing/flossing, preferably before bedtime.' if EN else 'Tagliches Nahrungserganzungsmittel. 1 Lutschtablette nach dem Zahneputzen langsam im Mund auflosen lassen.')),
        ('Route of Exposure' if EN else 'Expositionsroute', ('Oral: primary buccal/sublingual; secondary gastrointestinal.' if EN else 'Oral: primar bukkal/sublingual; sekundar gastrointestinal.')),
        ('Frequency' if EN else 'Haufigkeit', ('1 lozenge per day' if EN else '1 Lutschtablette pro Tag')),
    ]
    for lbl,val in prod_fields: field(doc,lbl,val)

    # ── 4. Available Documentation ──
    h1(doc, '4. '+('Available Documentation' if EN else 'Verfugbare Dokumentation'))
    p(doc, ('The following documentation was available or requested:' if EN else 'Folgende Dokumentation war verfugbar oder wurde angefordert:'))
    doc_h = ['Document Type','Status','Remarks'] if EN else ['Dokumententyp','Status','Anmerkungen']
    doc_r = [
        [('Full qualitative/quantitative composition' if EN else 'Vollstandige qualitative/quantitative Zusammensetzung'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
        [('Safety Data Sheets (SDS) — all raw materials' if EN else 'Sicherheitsdatenblatter (SDB) — alle Rohstoffe'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',('Per ingredient' if EN else 'Pro Inhaltsstoff')],
        [('Product specifications' if EN else 'Produktspezifikationen'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
        [('Certificates of Analysis (CoA)' if EN else 'Analysezertifikate (CoA)'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
        [('Microbiological test results' if EN else 'Mikrobiologische Testergebnisse'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
        [('Stability data' if EN else 'Stabilitatsdaten'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
        [('Product label / packaging' if EN else 'Produktetikett / Verpackung'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
        [('Published literature / clinical studies' if EN else 'Veroffentlichte Literatur / klinische Studien'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
        [('Strain identification certificates' if EN else 'Stammidentifizierungszertifikate'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
        [('Antibiotic resistance profiles' if EN else 'Antibiotikaresistenzprofile'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',('EFSA requirement' if EN else 'EFSA-Anforderung')],
        [('QPS status documentation' if EN else 'QPS-Statusdokumentation'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
        [('Heavy metals / contaminant testing' if EN else 'Schwermetall-/Kontaminantentests'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]','Pb, Cd, Hg, As'],
        [('Allergen risk assessment' if EN else 'Allergenbewertung'),'[Available/Pending]' if EN else '[Verfugbar/Ausstehend]',''],
    ]
    tbl(doc, doc_h, doc_r)

    # ── 5. Composition ──
    h1(doc, '5. '+('Qualitative and Quantitative Composition' if EN else 'Qualitative und quantitative Zusammensetzung'))
    comp_h = ['No.','Component' if EN else 'Bestandteil','CAS/EC','Function' if EN else 'Funktion','Amount' if EN else 'Menge','Remarks' if EN else 'Anmerkungen']
    comp_r = [
        ['1','Lactobacillus reuteri DSM 17938','N/A (biol.)',('Active — oral probiotic' if EN else 'Wirkstoff — orales Probiotikum'),'100M CFU' if EN else '100M KBE','EFSA QPS'],
        ['2','Lactobacillus paracasei [strain]','N/A (biol.)',('Active — oral probiotic' if EN else 'Wirkstoff — orales Probiotikum'),'[CFU]' if EN else '[KBE]','EFSA QPS'],
        ['3','Streptococcus salivarius K12','N/A (biol.)',('Active — oral probiotic' if EN else 'Wirkstoff — orales Probiotikum'),'1B CFU' if EN else '1 Mrd. KBE','EFSA QPS'],
        ['4','Streptococcus salivarius M18','N/A (biol.)',('Active — oral probiotic' if EN else 'Wirkstoff — orales Probiotikum'),'500M CFU' if EN else '500M KBE','EFSA QPS'],
        ['5','Xylitol','CAS 87-99-0 / EC 201-788-0',('Sweetener / dental agent' if EN else 'Sussmittel / Zahnpflegemittel'),'500 mg',('Laxative threshold ~20 g/day' if EN else 'Abfuhrende Schwelle ~20 g/Tag')],
        ['6','Peppermint oil (Mentha x piperita)','CAS 8006-90-4',('Flavouring' if EN else 'Aroma'),'q.s.','FEMA GRAS 2848'],
        ['7','[Bulking agent(s)]' if EN else '[Fullstoff(e)]','[CAS]',('Excipient' if EN else 'Hilfsstoff'),'[mg]','[e.g. isomalt]' if EN else '[z.B. Isomalt]'],
        ['8','[Binder(s)]' if EN else '[Bindemittel]','[CAS]',('Excipient' if EN else 'Hilfsstoff'),'[mg]',''],
        ['9','[Anti-caking agent(s)]' if EN else '[Trennmittel]','[CAS]',('Excipient' if EN else 'Hilfsstoff'),'[mg]',''],
    ]
    tbl(doc, comp_h, comp_r)
    p(doc, ('Note: Full excipient list and exact quantities must be provided by the commissioner.' if EN else 'Hinweis: Die vollstandige Hilfsstoffliste muss vom Auftraggeber bereitgestellt werden.'), italic=True)

    # ── 6. Methodology ──
    h1(doc, '6. '+('Methodology' if EN else 'Methodik'))
    if EN:
        p(doc,'This assessment follows established principles of toxicological risk assessment per:\n- WHO/IPCS Framework for Risk Assessment (2004)\n- EFSA Guidance on risk assessment of substances in food (2012)\n- EFSA Guidance on characterisation of microorganisms (2018)\n- EFSA Qualified Presumption of Safety (QPS) approach\n- Regulation (EC) No 178/2002, Directive 2002/46/EC\n- Relevant JECFA evaluations')
        p(doc,'The standard four-step paradigm is followed:\n1. Hazard Identification\n2. Hazard Characterisation (dose-response, NOAEL, LOAEL, ADI)\n3. Exposure Assessment\n4. Risk Characterisation (MoS/MoE calculation)')
        p(doc,'For probiotics, additionally: taxonomic identity, QPS status, absence of transferable AMR genes, history of safe use, clinical safety data, viability.')
    else:
        p(doc,'Diese Bewertung folgt etablierten Grundsatzen der toxikologischen Risikobewertung gemas:\n- WHO/IPCS-Rahmenwerk (2004)\n- EFSA-Leitlinie zur Risikobewertung von Stoffen in Lebensmitteln (2012)\n- EFSA-Leitlinie zur Charakterisierung von Mikroorganismen (2018)\n- EFSA QPS-Ansatz\n- Verordnung (EG) Nr. 178/2002, Richtlinie 2002/46/EG\n- Einschlagige JECFA-Bewertungen')
        p(doc,'Vierstufiges Paradigma:\n1. Gefahrenidentifizierung\n2. Gefahrencharakterisierung (Dosis-Wirkung, NOAEL, LOAEL, ADI)\n3. Expositionsabschatzung\n4. Risikocharakterisierung (MoS/MoE-Berechnung)')

    # ── 7. Hazard Identification ──
    h1(doc, '7. '+('Hazard Identification per Ingredient' if EN else 'Gefahrenidentifizierung pro Inhaltsstoff'))
    
    HL = [('id','Identification' if EN else 'Identifizierung'),
          ('acute','Acute Toxicity' if EN else 'Akute Toxizitat'),
          ('skin','Skin Irritation' if EN else 'Hautreizung'),
          ('sensitisation','Sensitisation' if EN else 'Sensibilisierung'),
          ('repeated','Repeated Dose Toxicity' if EN else 'Toxizitat bei wiederholter Dosis'),
          ('geno','Genotoxicity / Mutagenicity' if EN else 'Genotoxizitat / Mutagenitat'),
          ('carcino','Carcinogenicity' if EN else 'Kanzerogenitat'),
          ('repro','Reproductive Toxicity' if EN else 'Reproduktionstoxizitat'),
          ('critical','Critical Endpoint' if EN else 'Kritischer Endpunkt'),
          ('uncertainty','Uncertainties' if EN else 'Unsicherheiten')]

    # 7.1 L. reuteri
    hazard_block(doc, '7.1 Lactobacillus reuteri DSM 17938', {
        'id': ('Species: Lactobacillus reuteri | Strain: DSM 17938 | CAS: N/A (biological) | QPS: Listed, no safety concerns.' if EN else 'Spezies: Lactobacillus reuteri | Stamm: DSM 17938 | CAS: N/A (biologisch) | QPS: Gelistet, keine Sicherheitsbedenken.'),
        'acute': ('No acute toxicity reported. Studied at doses up to 10^10 CFU/day without adverse effects. Long history of safe use.' if EN else 'Keine akute Toxizitat berichtet. Untersucht bei Dosen bis 10^10 KBE/Tag ohne Nebenwirkungen. Lange Geschichte sicherer Verwendung.'),
        'skin': ('Not applicable (oral route).' if EN else 'Nicht zutreffend (orale Route).'),
        'sensitisation': ('No evidence. Verify allergen cross-contamination from culture medium.' if EN else 'Kein Nachweis. Allergen-Kreuzkontamination aus Kulturmedium prufen.'),
        'repeated': ('Long-term studies (up to 12 months) show no systemic adverse effects.' if EN else 'Langzeitstudien (bis 12 Monate) zeigen keine systemischen Nebenwirkungen.'),
        'geno': ('No evidence of genotoxicity. LAB do not produce genotoxic metabolites.' if EN else 'Kein Nachweis von Genotoxizitat. Milchsaurebakterien produzieren keine genotoxischen Metaboliten.'),
        'carcino': ('No evidence.' if EN else 'Kein Nachweis.'),
        'repro': ('Clinical studies in pregnant women show no adverse outcomes. Limited data.' if EN else 'Klinische Studien bei Schwangeren zeigen keine negativen Ergebnisse. Begrenzte Daten.'),
        'critical': ('No critical endpoint at intended dose. QPS status and clinical data support safety.' if EN else 'Kein kritischer Endpunkt bei beabsichtigter Dosis. QPS-Status und klinische Daten stutzen die Sicherheit.'),
        'uncertainty': ('Confirm strain-specific antibiotic resistance profile (EFSA). Verify viability at end of shelf life.' if EN else 'Stammspezifisches Antibiotikaresistenzprofil bestatigen (EFSA). Lebensfahigkeit am Ende der Haltbarkeit uberprufen.'),
    }, HL)

    # 7.2 L. paracasei
    hazard_block(doc, '7.2 Lactobacillus paracasei [Strain]', {
        'id': ('Species: Lactobacillus paracasei | Strain: [To be provided] | CAS: N/A | QPS: Listed.' if EN else 'Spezies: Lactobacillus paracasei | Stamm: [Anzugeben] | CAS: N/A | QPS: Gelistet.'),
        'acute': ('No acute toxicity at supplement-level doses. Safe use history in fermented dairy.' if EN else 'Keine akute Toxizitat bei Supplement-Dosen. Sichere Verwendung in fermentierten Milchprodukten.'),
        'skin': ('Not applicable (oral route).' if EN else 'Nicht zutreffend (orale Route).'),
        'sensitisation': ('No known potential. Assess culture media allergens.' if EN else 'Kein bekanntes Potenzial. Kulturmedien-Allergene bewerten.'),
        'repeated': ('Generally well tolerated. [Add strain-specific references.]' if EN else 'Allgemein gut vertraglich. [Stammspezifische Referenzen erganzen.]'),
        'geno': ('No evidence for the species.' if EN else 'Kein Nachweis fur die Spezies.'),
        'carcino': ('No evidence.' if EN else 'Kein Nachweis.'),
        'repro': ('Insufficient strain-specific data. No species-level concerns.' if EN else 'Unzureichende stammspezifische Daten. Keine Bedenken auf Speziesebene.'),
        'critical': ('No critical endpoint identified.' if EN else 'Kein kritischer Endpunkt identifiziert.'),
        'uncertainty': ('ACTION REQUIRED: Exact strain designation and antibiotic resistance profile needed.' if EN else 'MASSNAHME ERFORDERLICH: Genaue Stammbezeichnung und Antibiotikaresistenzprofil benotigt.'),
    }, HL)

    # 7.3 S. salivarius K12
    hazard_block(doc, '7.3 Streptococcus salivarius K12 (BLIS K12)', {
        'id': ('Species: S. salivarius | Strain: K12 | QPS: Listed. Predominant commensal of the human oral cavity.' if EN else 'Spezies: S. salivarius | Stamm: K12 | QPS: Gelistet. Vorherrschender Kommensal der menschlichen Mundhohle.'),
        'acute': ('No acute toxicity. Over 30 years of commercial use. Clinical doses up to 2.5x10^9 CFU/day safe.' if EN else 'Keine akute Toxizitat. Uber 30 Jahre kommerzielle Nutzung. Klinische Dosen bis 2,5x10^9 KBE/Tag sicher.'),
        'skin': ('Not applicable.' if EN else 'Nicht zutreffend.'),
        'sensitisation': ('No evidence. Normal human commensal.' if EN else 'Kein Nachweis. Normaler menschlicher Kommensal.'),
        'repeated': ('Long-term use (up to 6 months) shows no adverse effects. Transient colonisation.' if EN else 'Langzeitanwendung (bis 6 Monate) zeigt keine Nebenwirkungen. Vorubergehende Kolonisierung.'),
        'geno': ('No evidence.' if EN else 'Kein Nachweis.'),
        'carcino': ('No evidence.' if EN else 'Kein Nachweis.'),
        'repro': ('No specific data. No species-level concerns.' if EN else 'Keine spezifischen Daten. Keine Bedenken auf Speziesebene.'),
        'critical': ('No critical endpoint. QPS status + 30+ years commercial use.' if EN else 'Kein kritischer Endpunkt. QPS-Status + 30+ Jahre kommerzielle Nutzung.'),
        'uncertainty': ('Confirm antibiotic resistance profile from BLIS Technologies.' if EN else 'Antibiotikaresistenzprofil von BLIS Technologies bestatigen.'),
    }, HL)

    # 7.4 S. salivarius M18
    hazard_block(doc, '7.4 Streptococcus salivarius M18 (BLIS M18)', {
        'id': ('Species: S. salivarius | Strain: M18 | QPS: Listed.' if EN else 'Spezies: S. salivarius | Stamm: M18 | QPS: Gelistet.'),
        'acute': ('No acute toxicity observed in clinical trials.' if EN else 'Keine akute Toxizitat in klinischen Studien beobachtet.'),
        'skin': ('Not applicable.' if EN else 'Nicht zutreffend.'),
        'sensitisation': ('No evidence. Normal human oral commensal.' if EN else 'Kein Nachweis. Normaler oraler Kommensal.'),
        'repeated': ('Good tolerability with repeated daily use. No systemic effects.' if EN else 'Gute Vertraglichkeit bei taglicher Anwendung. Keine systemischen Effekte.'),
        'geno': ('No evidence.' if EN else 'Kein Nachweis.'),
        'carcino': ('No evidence.' if EN else 'Kein Nachweis.'),
        'repro': ('No specific data. No species-level concerns.' if EN else 'Keine spezifischen Daten. Keine Bedenken auf Speziesebene.'),
        'critical': ('No critical endpoint at intended dose.' if EN else 'Kein kritischer Endpunkt bei beabsichtigter Dosis.'),
        'uncertainty': ('Confirm antibiotic resistance profile. Document viability specs.' if EN else 'Antibiotikaresistenzprofil bestatigen. Lebensfahigkeitsspezifikationen dokumentieren.'),
    }, HL)

    # 7.5 Xylitol
    hazard_block(doc, '7.5 Xylitol', {
        'id': ('Xylitol (D-xylitol) | CAS: 87-99-0 | EC: 201-788-0 | E 967 | MW: 152.15 g/mol | JECFA ADI: not specified (safe at normal dietary levels). Approved food additive (quantum satis) under Reg (EC) 1333/2008.' if EN else 'Xylitol (D-Xylitol) | CAS: 87-99-0 | EC: 201-788-0 | E 967 | MW: 152,15 g/mol | JECFA ADI: nicht festgelegt. Zugelassener Lebensmittelzusatzstoff gemas VO (EG) 1333/2008.'),
        'acute': ('Oral LD50 (rat): >20 g/kg bw. Very low acute toxicity. High single doses (>20 g) may cause transient GI discomfort.' if EN else 'Orale LD50 (Ratte): >20 g/kg KG. Sehr geringe akute Toxizitat. Hohe Einzeldosen (>20 g) konnen vorubergehende GI-Beschwerden verursachen.'),
        'skin': ('Not classified as irritant.' if EN else 'Nicht als reizend eingestuft.'),
        'sensitisation': ('Not classified. No reports of allergic reactions.' if EN else 'Nicht eingestuft. Keine Berichte uber allergische Reaktionen.'),
        'repeated': ('2-year rat study: NOAEL 5 g/kg bw/day. Caecal enlargement (adaptive, not adverse). JECFA considers GI effects pharmacological, not toxicological.' if EN else '2-Jahres-Rattenstudie: NOAEL 5 g/kg KG/Tag. Zakalerweiterung (adaptiv, nicht nachteilig). JECFA betrachtet GI-Effekte als pharmakologisch.'),
        'geno': ('Negative in Ames, chromosomal aberration, micronucleus assays. No genotoxic potential.' if EN else 'Negativ in Ames-, Chromosomenaberrations-, Mikronukleus-Tests. Kein genotoxisches Potenzial.'),
        'carcino': ('No evidence. Not classified by IARC.' if EN else 'Kein Nachweis. Nicht von IARC eingestuft.'),
        'repro': ('No evidence at nutritionally relevant doses.' if EN else 'Kein Nachweis bei ernahrungsrelevanten Dosen.'),
        'critical': ('GI tolerance is dose-limiting. Laxative threshold ~20-30 g/day adults. At 500 mg/day: well below threshold (2.5%).' if EN else 'GI-Toleranz ist dosislimitierend. Abfuhrende Schwelle ~20-30 g/Tag. Bei 500 mg/Tag: weit unter Schwelle (2,5%).'),
        'uncertainty': ('Low. Robust safety dossier. Note: highly toxic to dogs — relevant for labelling.' if EN else 'Gering. Robustes Sicherheitsdossier. Hinweis: hochgiftig fur Hunde — relevant fur Kennzeichnung.'),
    }, HL)

    # 7.6 Peppermint
    hazard_block(doc, '7.6 '+('Natural Peppermint Flavouring' if EN else 'Naturliches Pfefferminzaroma'), {
        'id': ('Peppermint oil (Mentha x piperita L.) | CAS: 8006-90-4 | FEMA GRAS 2848 | Primary constituent: menthol (30-55%). GRAS as flavouring. Listed in Reg (EC) 1334/2008.' if EN else 'Pfefferminzol (Mentha x piperita L.) | CAS: 8006-90-4 | FEMA GRAS 2848 | Hauptbestandteil: Menthol (30-55%). GRAS als Aroma. Gelistet in VO (EG) 1334/2008.'),
        'acute': ('Oral LD50 (rat): ~4.4 g/kg bw. No concern at flavouring levels.' if EN else 'Orale LD50 (Ratte): ~4,4 g/kg KG. Kein Bedenken bei Aromamengen.'),
        'skin': ('Mild irritant at high conc. Not relevant for oral lozenge.' if EN else 'Leicht reizend bei hoher Konz. Nicht relevant fur orale Lutschtablette.'),
        'sensitisation': ('Rare contact sensitisation to menthol reported. Oral exposure at flavouring levels not associated.' if EN else 'Seltene Kontaktsensibilisierung gegen Menthol berichtet. Orale Exposition bei Aromamengen nicht assoziiert.'),
        'repeated': ('Menthol ADI (JECFA): 0-4 mg/kg bw/day. Flavouring level exposure expected well below.' if EN else 'Menthol ADI (JECFA): 0-4 mg/kg KG/Tag. Exposition bei Aromamenge weit darunter.'),
        'geno': ('Menthol: negative in standard battery. No concern at flavouring concentrations.' if EN else 'Menthol: negativ im Standardtestprogramm. Kein Bedenken bei Aromamengen.'),
        'carcino': ('No evidence.' if EN else 'Kein Nachweis.'),
        'repro': ('High-dose animal effects only. Not relevant at flavouring levels.' if EN else 'Nur Hochdosis-Tiereffekte. Nicht relevant bei Aromamengen.'),
        'critical': ('No concern at flavouring concentration.' if EN else 'Kein Bedenken bei Aromamenge.'),
        'uncertainty': ('ACTION REQUIRED: Exact quantity per lozenge needed for quantitative assessment.' if EN else 'MASSNAHME ERFORDERLICH: Genaue Menge pro Lutschtablette fur quantitative Bewertung benotigt.'),
    }, HL)

    # 7.7 Excipients
    h3(doc, '7.7 '+('Excipients' if EN else 'Hilfsstoffe'))
    if EN:
        p(doc, 'The exact excipients have not been fully disclosed. Common excipients in probiotic lozenges include: isomalt (E 953), mannitol (E 421), sorbitol (E 420), microcrystalline cellulose (E 460i), magnesium stearate (E 470b), silicon dioxide (E 551). All are approved food additives with established safety profiles.')
        p(doc, 'ACTION REQUIRED: The commissioner must provide the complete excipient list with exact quantities for this section to be finalised.', bold=True)
    else:
        p(doc, 'Die genauen Hilfsstoffe wurden nicht vollstandig offengelegt. Ubliche Hilfsstoffe in probiotischen Lutschtabletten: Isomalt (E 953), Mannitol (E 421), Sorbitol (E 420), Mikrokristalline Cellulose (E 460i), Magnesiumstearat (E 470b), Siliciumdioxid (E 551). Alle sind zugelassene Lebensmittelzusatzstoffe.')
        p(doc, 'MASSNAHME ERFORDERLICH: Der Auftraggeber muss die vollstandige Hilfsstoffliste mit genauen Mengen bereitstellen.', bold=True)

    # ── 8. Exposure Assessment ──
    h1(doc, '8. '+('Exposure Assessment' if EN else 'Expositionsabschatzung'))
    exp_fields = [
        ('Use Scenario' if EN else 'Verwendungsszenario', ('Daily oral care supplement. 1 lozenge dissolved slowly after brushing/flossing, ~5-10 min dissolution. No eating/drinking for 30 min.' if EN else 'Tagliches Mundpflegeerganzungsmittel. 1 Lutschtablette nach Zahneputzen langsam auflosen (~5-10 Min). 30 Min nicht essen/trinken.')),
        ('Frequency' if EN else 'Haufigkeit', ('Once daily (1 lozenge/day)' if EN else 'Einmal taglich (1 Lutschtablette/Tag)')),
        ('Duration' if EN else 'Dauer', ('Long-term daily use (continuous, subscription model 30-day cycles)' if EN else 'Langfristige tagliche Anwendung (kontinuierlich, Abo-Modell 30-Tage-Zyklen)')),
        ('Target Population' if EN else 'Zielpopulation', ('Healthy adults >= 18 years. BW assumption: 60 kg (conservative).' if EN else 'Gesunde Erwachsene >= 18 Jahre. KG-Annahme: 60 kg (konservativ).')),
        ('Route' if EN else 'Route', ('Primary: buccal/sublingual. Secondary: GI (after swallowing). Dermal/inhalation: N/A.' if EN else 'Primar: bukkal/sublingual. Sekundar: GI (nach Verschlucken). Dermal/Inhalation: N/A.')),
    ]
    for lbl,val in exp_fields: field(doc,lbl,val)

    h2(doc, ('Estimated Daily Exposure' if EN else 'Geschatzte tagliche Exposition'))
    exp_h = [('Ingredient' if EN else 'Inhaltsstoff'),'Amount/Dose' if EN else 'Menge/Dosis','mg/kg bw (60 kg)',('Reference Value' if EN else 'Referenzwert'),('Assessment' if EN else 'Bewertung')]
    exp_r = [
        ['L. reuteri DSM 17938','100M CFU','N/A',('No ADI for QPS probiotics' if EN else 'Kein ADI fur QPS-Probiotika'),('Acceptable' if EN else 'Akzeptabel')],
        ['L. paracasei','[CFU]','N/A','[Pending]' if EN else '[Ausstehend]','[TBD]'],
        ['S. salivarius K12','1B CFU','N/A',('No ADI' if EN else 'Kein ADI'),('Acceptable' if EN else 'Akzeptabel')],
        ['S. salivarius M18','500M CFU','N/A',('No ADI' if EN else 'Kein ADI'),('Acceptable' if EN else 'Akzeptabel')],
        ['Xylitol','500 mg','8.3 mg/kg',('JECFA ADI: not specified; laxative ~20g/day' if EN else 'JECFA ADI: nicht festgelegt; abfuhrend ~20g/Tag'),('Acceptable (2.5% of threshold)' if EN else 'Akzeptabel (2,5% der Schwelle)')],
        [('Peppermint oil' if EN else 'Pfefferminzol'),'q.s.','[TBD]',('Menthol ADI: 0-4 mg/kg bw/day' if EN else 'Menthol ADI: 0-4 mg/kg KG/Tag'),('[Requires quantification]' if EN else '[Quantifizierung erforderlich]')],
    ]
    tbl(doc, exp_h, exp_r)

    # ── 9. Risk Characterisation ──
    h1(doc, '9. '+('Risk Characterisation' if EN else 'Risikocharakterisierung'))
    if EN:
        p(doc, 'Based on available data:\n- The four probiotic strains are QPS-listed and used at or below clinically studied doses.\n- Xylitol at 500 mg/day = 2.5% of laxative threshold. Negligible risk.\n- Peppermint at flavouring level: no concern expected. Quantitative confirmation required.\n- Excipients: pending full composition.\n\nNo unacceptable risks identified for assessed components. Overall safety profile is favourable, pending data gap completion.')
        p(doc, 'Conditions:\n- Valid only for described formulation and use scenario.\n- Assumes compliance with recommended dose (1 lozenge/day).\n- Does not cover drug interactions or immunocompromised individuals.\n- Any formulation change requires reassessment.')
    else:
        p(doc, 'Basierend auf verfugbaren Daten:\n- Die vier probiotischen Stamme sind QPS-gelistet und werden in oder unter klinisch untersuchten Dosen verwendet.\n- Xylitol bei 500 mg/Tag = 2,5% der abfuhrenden Schwelle. Vernachlassigbares Risiko.\n- Pfefferminz als Aroma: kein Bedenken erwartet. Quantitative Bestatigung erforderlich.\n- Hilfsstoffe: ausstehend.\n\nKeine inakzeptablen Risiken fur bewertete Bestandteile identifiziert. Gesamtsicherheitsprofil gunstig, vorbehaltlich Schliessung der Datenlucken.')
        p(doc, 'Bedingungen:\n- Gultig nur fur beschriebene Formulierung und Verwendungsszenario.\n- Setzt Einhaltung der empfohlenen Dosis voraus (1 Lutschtablette/Tag).\n- Deckt keine Arzneimittelwechselwirkungen oder immungeschwachte Personen ab.\n- Jede Formulierungsanderung erfordert Neubewertung.')

    # ── 10. Uncertainties ──
    h1(doc, '10. '+('Uncertainties and Limitations' if EN else 'Unsicherheiten und Einschrankungen'))
    if EN:
        items = ['Incomplete composition data — full excipient list not provided.','L. paracasei strain designation not disclosed. Strain-level safety data required per EFSA.','Peppermint oil quantity not specified — quantitative assessment pending.','Antibiotic resistance profiles for all four strains required (EFSA).','Stability and viability data over shelf life needed.','Heavy metal testing (Pb, Cd, Hg, As) and microbiological quality needed.','Aggregate xylitol/menthol exposure from other dietary sources not assessed.','Assessment limited to healthy adults. Pregnant/breastfeeding, immunocompromised, and children not fully evaluated.','Accuracy depends on commissioner-provided data.','Periodic review required when new data/regulatory guidance available.']
    else:
        items = ['Unvollstandige Zusammensetzungsdaten — vollstandige Hilfsstoffliste nicht bereitgestellt.','L. paracasei-Stammbezeichnung nicht offengelegt. Stammspezifische Sicherheitsdaten gemas EFSA erforderlich.','Pfefferminzolmenge nicht angegeben — quantitative Bewertung ausstehend.','Antibiotikaresistenzprofile fur alle vier Stamme erforderlich (EFSA).','Stabilitits- und Lebensfihigkeitsdaten uber die Haltbarkeit benotigt.','Schwermetalltests (Pb, Cd, Hg, As) und mikrobiologische Qualitat benotigt.','Aggregierte Xylitol/Menthol-Exposition aus anderen Quellen nicht bewertet.','Bewertung auf gesunde Erwachsene beschrankt.','Richtigkeit hangt von Auftraggeber-Daten ab.','Periodische Uberprufung bei neuen Daten/regulatorischen Anderungen erforderlich.']
    for i,item in enumerate(items,1):
        p(doc, f'{i}. {item}')

    # ── 11. Conclusion ──
    h1(doc, '11. '+('Conclusion' if EN else 'Schlussfolgerung'))
    if EN:
        p(doc, 'Based on the information available, and subject to the uncertainties in Section 10:')
        p(doc, '1. The probiotic strains (L. reuteri DSM 17938, L. paracasei, S. salivarius K12, S. salivarius M18) belong to EFSA QPS-listed species. No toxicological concerns at intended daily doses.')
        p(doc, '2. Xylitol at 500 mg/day is well within the range of safe intake for adults.')
        p(doc, '3. Natural peppermint flavouring at customary levels is not expected to present a safety concern.')
        p(doc, '4. The Product, under intended conditions of use (1 lozenge daily, dissolved in mouth, by healthy adults), is considered toxicologically acceptable, provided that:\n   a) Complete quantitative formulation is confirmed;\n   b) Antibiotic resistance profiles are verified;\n   c) L. paracasei strain designation is provided;\n   d) Finished product meets quality/purity specifications.')
        p(doc, '5. This conclusion applies exclusively to the described formulation and use scenario. Any modification requires reassessment.')
        p(doc, '6. The product label warnings (children, pets, pregnancy, xylitol laxative effect) are appropriate and sufficient.')
        h2(doc, 'Compact Conclusion')
        p(doc, 'OralBiome Pro, containing four QPS-listed probiotic strains, xylitol, and natural peppermint flavouring in a dissolvable lozenge, is considered toxicologically acceptable for daily use by healthy adults at the recommended dose of one lozenge per day, subject to completion of the outstanding data requirements identified in this report.', italic=True)
    else:
        p(doc, 'Basierend auf den verfugbaren Informationen und vorbehaltlich der Unsicherheiten in Abschnitt 10:')
        p(doc, '1. Die probiotischen Stamme (L. reuteri DSM 17938, L. paracasei, S. salivarius K12, S. salivarius M18) gehoren zu EFSA QPS-gelisteten Spezies. Keine toxikologischen Bedenken bei beabsichtigten Tagesdosen.')
        p(doc, '2. Xylitol bei 500 mg/Tag liegt gut im sicheren Aufnahmebereich fur Erwachsene.')
        p(doc, '3. Naturliches Pfefferminzaroma in ublichen Mengen stellt kein Sicherheitsbedenken dar.')
        p(doc, '4. Das Produkt wird unter bestimmungsgemassen Bedingungen (1 Lutschtablette taglich, im Mund aufgelost, von gesunden Erwachsenen) als toxikologisch akzeptabel erachtet, vorausgesetzt:\n   a) Vollstandige Zusammensetzung bestatigt;\n   b) Antibiotikaresistenzprofile verifiziert;\n   c) L. paracasei-Stammbezeichnung bereitgestellt;\n   d) Fertigprodukt erfullt Qualitats-/Reinheitsspezifikationen.')
        p(doc, '5. Diese Schlussfolgerung gilt ausschliesslich fur die beschriebene Formulierung. Jede Anderung erfordert Neubewertung.')
        h2(doc, 'Kompakte Schlussfolgerung')
        p(doc, 'OralBiome Pro, mit vier QPS-gelisteten probiotischen Stammen, Xylitol und naturlichem Pfefferminzaroma als auflosbare Lutschtablette, wird als toxikologisch akzeptabel fur die tagliche Anwendung durch gesunde Erwachsene bei der empfohlenen Dosis von einer Lutschtablette pro Tag erachtet, vorbehaltlich der Erfullung der in diesem Bericht identifizierten ausstehenden Datenanforderungen.', italic=True)

    # ── 12. Recommendations ──
    h1(doc, '12. '+('Recommendations' if EN else 'Empfehlungen'))
    if EN:
        recs = ['Provide full quantitative formulation (all excipients + exact quantities).','Disclose L. paracasei exact strain designation + strain identity certificates.','Provide antibiotic susceptibility testing per EFSA guidelines for all 4 strains.','Provide peppermint oil specification (quantity/lozenge, menthol content, quality standard).','Provide accelerated and real-time stability data.','Provide heavy metal (Pb, Cd, Hg, As) and microbiological finished product testing.','Confirm allergen status of all raw materials (including culture media cross-contamination).','Review labelling for completeness (ingredient list, CFU declaration, warnings, storage, batch, best-before).','Establish post-market adverse event monitoring system.','Schedule reassessment every 2-3 years or upon formulation/regulatory changes.']
    else:
        recs = ['Vollstandige quantitative Formulierung bereitstellen (alle Hilfsstoffe + genaue Mengen).','L. paracasei-Stammbezeichnung offenlegen + Stammidentifikatszertifikate.','Antibiotikaempfindlichkeitstests gemas EFSA-Richtlinien fur alle 4 Stamme.','Pfefferminzol-Spezifikation bereitstellen (Menge/Lutschtablette, Mentholgehalt).','Beschleunigte und Echtzeit-Stabilitatsdaten bereitstellen.','Schwermetall- (Pb, Cd, Hg, As) und mikrobiologische Fertigprodukttests.','Allergenstatus aller Rohstoffe bestatigen.','Kennzeichnung auf Vollstandigkeit prufen.','Postmarkt-Nebenwirkungs-Uberwachungssystem einrichten.','Neubewertung alle 2-3 Jahre oder bei Formulierungs-/Regulierungsanderungen.']
    for i,r in enumerate(recs,1): p(doc, f'{i}. {r}')

    # ── 13. Declaration ──
    h1(doc, '13. '+('Declaration of Competence and Independence' if EN else 'Erklarung der Sachkunde und Unabhangigkeit'))
    if EN:
        p(doc, 'I, the undersigned, declare that this toxicological safety assessment has been prepared to the best of my scientific knowledge and professional judgement, in accordance with generally accepted principles of toxicology and risk assessment.')
        p(doc, 'I further declare that I have no financial interest in, or other conflict of interest with, the commissioning party (CALQIX B.V.) or its products, other than the agreed professional fee.')
    else:
        p(doc, 'Ich, der/die Unterzeichnende, erklare, dass diese toxikologische Sicherheitsbewertung nach bestem wissenschaftlichen Wissen und beruflichen Urteilsvermogen erstellt wurde, in Ubereinstimmung mit allgemein anerkannten Grundsatzen der Toxikologie und Risikobewertung.')
        p(doc, 'Ich erklare ferner, dass ich kein finanzielles Interesse an oder sonstigen Interessenkonflikt mit dem Auftraggeber (CALQIX B.V.) oder seinen Produkten habe, ausser dem vereinbarten Berufshonorar.')
    
    field(doc, 'Name' if EN else 'Name', '[Full name]' if EN else '[Vollstandiger Name]')
    field(doc, 'Title' if EN else 'Titel', '[e.g. Registered Toxicologist (ERT)]')
    field(doc, 'Date' if EN else 'Datum', '[DD Month YYYY]' if EN else '[TT Monat JJJJ]')
    field(doc, 'Signature' if EN else 'Unterschrift', '________________________________________')

    # ── 14. Annexes ──
    h1(doc, '14. '+('Annexes' if EN else 'Anhange'))
    if EN:
        annexes = ['Annex A — Full qualitative and quantitative composition','Annex B — Safety Data Sheets (SDS)','Annex C — Certificates of Analysis','Annex D — Strain identification certificates','Annex E — Antibiotic resistance test results','Annex F — Stability study reports','Annex G — Heavy metal and contaminant results','Annex H — Product label / packaging artwork','Annex I — Published literature and clinical summaries','Annex J — Exposure calculation worksheets','Annex K — Allergen risk assessment','Annex L — HACCP plan summary']
    else:
        annexes = ['Anhang A — Vollstandige Zusammensetzung','Anhang B — Sicherheitsdatenblatter (SDB)','Anhang C — Analysezertifikate','Anhang D — Stammidentifizierungszertifikate','Anhang E — Antibiotikaresistenz-Testergebnisse','Anhang F — Stabilitatsstudienberichte','Anhang G — Schwermetall- und Kontaminantenergebnisse','Anhang H — Produktetikett / Verpackung','Anhang I — Literatur und klinische Zusammenfassungen','Anhang J — Expositionsberechnungen','Anhang K — Allergenbewertung','Anhang L — HACCP-Planzusammenfassung']
    for a in annexes: p(doc, a)

    # ── Optional Extras ──
    h1(doc, ('Optional Additional Sections' if EN else 'Optionale zusatzliche Abschnitte'))
    opts = [
        'Executive Summary','Version History' if EN else 'Versionshistorie',
        'Confidentiality Statement' if EN else 'Vertraulichkeitserklarung',
        'Regulatory Framework' if EN else 'Regulatorischer Rahmen',
        'References' if EN else 'Literaturverzeichnis',
        'Glossary' if EN else 'Glossar',
        'EFSA QPS List Extract' if EN else 'EFSA QPS-Listenauszug'
    ]
    for o in opts: p(doc, f'- {o}')

    # ── Validation Note ──
    doc.add_page_break()
    h1(doc, ('Validation Note' if EN else 'Validierungshinweis'))
    if EN:
        p(doc, f'This document has been prepared as a professional template pre-populated with product-specific data for CALQIX OralBiome Pro. It is intended as a working document for use by a qualified toxicologist. The final assessment must be reviewed, completed, validated, and signed by a suitably qualified professional before it may be relied upon for regulatory compliance.\n\nTemplate version: 1.0\nGenerated: {today}\nPrepared by: Cascade AI (template generation only; no toxicological validation)\nStatus: DRAFT — Requires expert review and completion', italic=True)
    else:
        p(doc, f'Dieses Dokument wurde als professionelle Vorlage mit produktspezifischen Daten fur CALQIX OralBiome Pro erstellt. Es ist als Arbeitsdokument fur einen qualifizierten Toxikologen gedacht. Die endgultige Bewertung muss von einem entsprechend qualifizierten Fachmann uberpruft, vervollstandigt, validiert und unterzeichnet werden.\n\nVorlagenversion: 1.0\nErstellt: {today}\nErstellt von: Cascade AI (nur Vorlagenerstellung; keine toxikologische Validierung)\nStatus: ENTWURF — Erfordert Expertenprufung und Vervollstandigung', italic=True)

    return doc


if __name__ == '__main__':
    en_doc = build('en')
    de_doc = build('de')
    en_path = os.path.join(OUT_DIR, 'Toxicological_Report_OralBiome_Pro_EN.docx')
    de_path = os.path.join(OUT_DIR, 'Toxicological_Report_OralBiome_Pro_DE.docx')
    en_doc.save(en_path)
    de_doc.save(de_path)
    print(f'EN: {en_path}')
    print(f'DE: {de_path}')
